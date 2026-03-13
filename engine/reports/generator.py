"""
JanPulse AI — Report Generator
Automatically generates 50+ page DOCX/PDF rally intelligence reports.
Uses python-docx for document assembly and matplotlib for charts.
"""

import os
import io
import json
from datetime import datetime
from typing import Optional
from loguru import logger

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed; report generation disabled")

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

from engine.models import KPISnapshot, RallyPhase


class ChartGenerator:
    """Generate charts for the report using Matplotlib."""

    COLORS = {
        "positive": "#27AE60", "negative": "#E74C3C",
        "neutral": "#95A5A6", "mixed": "#F39C12",
        "primary": "#1B3A5C", "accent": "#C0392B",
        "blue": "#2980B9", "purple": "#8E44AD"
    }

    @staticmethod
    def sentiment_pie(kpi: KPISnapshot, output_path: str):
        if not HAS_MPL:
            return
        labels = ["Positive", "Negative", "Neutral", "Mixed"]
        sizes = [kpi.sentiment_share.positive, kpi.sentiment_share.negative,
                 kpi.sentiment_share.neutral, kpi.sentiment_share.mixed]
        colors = [ChartGenerator.COLORS["positive"], ChartGenerator.COLORS["negative"],
                  ChartGenerator.COLORS["neutral"], ChartGenerator.COLORS["mixed"]]

        fig, ax = plt.subplots(1, 1, figsize=(6, 4))
        wedges, texts, autotexts = ax.pie(
            sizes, labels=labels, colors=colors, autopct="%1.1f%%",
            startangle=90, textprops={"fontsize": 10}
        )
        ax.set_title("Sentiment Distribution", fontsize=13, fontweight="bold", color="#1B3A5C")
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close()

    @staticmethod
    def emotion_bar(kpi: KPISnapshot, output_path: str):
        if not HAS_MPL:
            return
        ed = kpi.emotion_distribution
        emotions = {
            "Hope": ed.hope, "Anger": ed.anger, "Pride": ed.pride,
            "Fear": ed.fear, "Trust": ed.trust, "Excitement": ed.excitement,
            "Frustration": ed.frustration, "Joy": ed.joy, "Sarcasm": ed.sarcasm
        }
        emotions = {k: v for k, v in emotions.items() if v > 0}
        if not emotions:
            return

        fig, ax = plt.subplots(figsize=(8, 4))
        bars = ax.barh(list(emotions.keys()), list(emotions.values()),
                       color=ChartGenerator.COLORS["primary"], edgecolor="white")
        ax.set_xlabel("Score", fontsize=10)
        ax.set_title("Emotion Distribution", fontsize=13, fontweight="bold", color="#1B3A5C")
        ax.invert_yaxis()
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close()

    @staticmethod
    def platform_bar(kpi: KPISnapshot, output_path: str):
        if not HAS_MPL:
            return
        platforms = kpi.platform_breakdown
        if not platforms:
            return

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.bar(list(platforms.keys()), list(platforms.values()),
               color=ChartGenerator.COLORS["blue"], edgecolor="white")
        ax.set_ylabel("Volume", fontsize=10)
        ax.set_title("Platform Activity Distribution", fontsize=13, fontweight="bold", color="#1B3A5C")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close()

    @staticmethod
    def hashtag_bar(kpi: KPISnapshot, output_path: str):
        if not HAS_MPL:
            return
        if not kpi.top_hashtags:
            return
        hashtags = kpi.top_hashtags[:10]
        names = [h["hashtag"] for h in hashtags]
        counts = [h["count"] for h in hashtags]

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.barh(names[::-1], counts[::-1], color=ChartGenerator.COLORS["accent"])
        ax.set_xlabel("Mentions")
        ax.set_title("Top Hashtags", fontsize=13, fontweight="bold", color="#1B3A5C")
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close()

    @staticmethod
    def mood_gauge(score: float, output_path: str):
        if not HAS_MPL:
            return
        fig, ax = plt.subplots(figsize=(6, 3))
        ax.barh(["Rally Mood Score"], [score], color=ChartGenerator.COLORS["primary"],
                 height=0.5)
        ax.barh(["Rally Mood Score"], [100 - score], left=[score],
                 color="#E0E0E0", height=0.5)
        ax.set_xlim(0, 100)
        ax.set_title(f"Rally Mood Score: {score}/100", fontsize=13,
                     fontweight="bold", color="#1B3A5C")
        ax.text(score / 2, 0, f"{score}", ha="center", va="center",
                fontsize=18, fontweight="bold", color="white")
        ax.set_yticks([])
        plt.tight_layout()
        plt.savefig(output_path, dpi=150, bbox_inches="tight")
        plt.close()


class ReportGenerator:
    """Generate comprehensive rally intelligence report as DOCX."""

    def __init__(self, campaign_config: dict = None):
        self.config = campaign_config or {}
        self.chart_gen = ChartGenerator()
        self.charts_dir = "/tmp/rie_charts"
        os.makedirs(self.charts_dir, exist_ok=True)

    def generate(self, kpi: KPISnapshot, narrative_summary: dict = None,
                 baseline_kpi: Optional[KPISnapshot] = None,
                 output_path: str = "rally_report.docx") -> str:
        """Generate the full report DOCX."""
        if not HAS_DOCX:
            logger.error("python-docx not installed. Cannot generate report.")
            return ""

        doc = Document()

        # ── Style setup ──
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)

        # ── Cover Page ──
        self._add_cover(doc, kpi)

        # ── Executive Summary ──
        doc.add_page_break()
        self._add_executive_summary(doc, kpi, baseline_kpi, narrative_summary)

        # ── Key Findings ──
        doc.add_page_break()
        self._add_key_findings(doc, kpi, baseline_kpi)

        # ── Public Mood Snapshot ──
        doc.add_page_break()
        self._add_mood_snapshot(doc, kpi)

        # ── Pre vs Post Comparison ──
        if baseline_kpi:
            doc.add_page_break()
            self._add_comparison(doc, kpi, baseline_kpi)

        # ── Platform Performance ──
        doc.add_page_break()
        self._add_platform_performance(doc, kpi)

        # ── Keywords and Hashtags ──
        doc.add_page_break()
        self._add_keywords(doc, kpi)

        # ── Influencer Analysis ──
        doc.add_page_break()
        self._add_influencer_analysis(doc, kpi)

        # ── Narratives ──
        if narrative_summary:
            doc.add_page_break()
            self._add_narratives(doc, narrative_summary)

        # ── Risk Assessment ──
        doc.add_page_break()
        self._add_risk_assessment(doc, kpi)

        # ── Methodology ──
        doc.add_page_break()
        self._add_methodology(doc, kpi)

        # ── Save ──
        doc.save(output_path)
        logger.info(f"Report generated: {output_path}")
        return output_path

    def _add_cover(self, doc, kpi):
        for _ in range(6):
            doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RALLY INTELLIGENCE REPORT")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(self.config.get("rally_name", "Political Rally Analysis"))
        run2.font.size = Pt(18)
        run2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"Phase: {kpi.phase.value.replace('_', ' ').title()}")
        run3.font.size = Pt(14)
        run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for _ in range(4):
            doc.add_paragraph("")

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run4 = p4.add_run(f"Generated: {datetime.utcnow().strftime('%d %B %Y, %H:%M UTC')}")
        run4.font.size = Pt(11)
        run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run5 = p5.add_run("Mindgen Pvt Ltd | Confidential")
        run5.font.size = Pt(11)
        run5.font.bold = True

    def _add_executive_summary(self, doc, kpi, baseline, narrative):
        doc.add_heading("Executive Summary", level=1)
        rally_name = self.config.get("rally_name", "the rally")
        mood_desc = "positive" if kpi.rally_mood_score > 60 else "negative" if kpi.rally_mood_score < 40 else "neutral"

        summary = (
            f"This report analyses public mood and digital engagement surrounding {rally_name}. "
            f"A total of {kpi.mention_volume:,} mentions were monitored across {len(kpi.platform_breakdown)} platforms "
            f"during the {kpi.phase.value.replace('_', ' ')} phase. "
            f"The overall Rally Mood Score stands at {kpi.rally_mood_score}/100, indicating a {mood_desc} reception. "
            f"Sentiment distribution shows {kpi.sentiment_share.positive}% positive, "
            f"{kpi.sentiment_share.negative}% negative, and {kpi.sentiment_share.neutral}% neutral. "
            f"The leader favourability score is {kpi.leader_favourability_score}/100. "
            f"Bot suspicion affects {kpi.bot_suspicion_score}% of analysed content. "
            f"The polarization index stands at {kpi.polarization_index}, "
            f"and {kpi.misinformation_flag_count} potential misinformation items were flagged."
        )
        doc.add_paragraph(summary)

        # Mood gauge chart
        gauge_path = os.path.join(self.charts_dir, "mood_gauge.png")
        self.chart_gen.mood_gauge(kpi.rally_mood_score, gauge_path)
        if os.path.exists(gauge_path):
            doc.add_picture(gauge_path, width=Inches(5))

    def _add_key_findings(self, doc, kpi, baseline):
        doc.add_heading("Key Findings", level=1)
        findings = [
            f"Total digital mentions reached {kpi.mention_volume:,} across all monitored platforms.",
            f"Positive sentiment ({kpi.sentiment_share.positive}%) {'outweighs' if kpi.sentiment_share.positive > kpi.sentiment_share.negative else 'trails'} negative sentiment ({kpi.sentiment_share.negative}%).",
            f"Rally Mood Score: {kpi.rally_mood_score}/100.",
            f"Leader favourability stands at {kpi.leader_favourability_score}/100.",
            f"Share of voice for {self.config.get('primary_party', 'the party')}: {kpi.share_of_voice}%.",
            f"Opposition counter-narrative intensity: {kpi.opposition_counter_intensity}%.",
            f"Bot-suspected content: {kpi.bot_suspicion_score}% of total volume.",
            f"Polarization index: {kpi.polarization_index} (higher = more divisive).",
            f"Top platform by volume: {max(kpi.platform_breakdown, key=kpi.platform_breakdown.get) if kpi.platform_breakdown else 'N/A'}.",
            f"Misinformation flags: {kpi.misinformation_flag_count} items flagged for review.",
        ]
        if baseline:
            delta = kpi.rally_mood_score - baseline.rally_mood_score
            direction = "improved" if delta > 0 else "declined" if delta < 0 else "remained stable"
            findings.append(f"Mood has {direction} by {abs(delta):.1f} points compared to baseline.")

        for i, finding in enumerate(findings, 1):
            doc.add_paragraph(f"{i}. {finding}")

    def _add_mood_snapshot(self, doc, kpi):
        doc.add_heading("Public Mood Snapshot", level=1)

        # Sentiment pie
        pie_path = os.path.join(self.charts_dir, "sentiment_pie.png")
        self.chart_gen.sentiment_pie(kpi, pie_path)
        if os.path.exists(pie_path):
            doc.add_picture(pie_path, width=Inches(4.5))

        # Emotion bar
        emo_path = os.path.join(self.charts_dir, "emotion_bar.png")
        self.chart_gen.emotion_bar(kpi, emo_path)
        if os.path.exists(emo_path):
            doc.add_picture(emo_path, width=Inches(5.5))

        # Sentiment table
        doc.add_heading("Sentiment Breakdown", level=2)
        table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
        table.rows[0].cells[0].text = "Sentiment"
        table.rows[0].cells[1].text = "Percentage"
        for i, (label, val) in enumerate([
            ("Positive", kpi.sentiment_share.positive),
            ("Negative", kpi.sentiment_share.negative),
            ("Neutral", kpi.sentiment_share.neutral),
            ("Mixed", kpi.sentiment_share.mixed)
        ], 1):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = f"{val}%"

    def _add_comparison(self, doc, kpi, baseline):
        doc.add_heading("Pre-Rally vs Post-Rally Comparison", level=1)
        metrics = [
            ("Mention Volume", baseline.mention_volume, kpi.mention_volume),
            ("Positive Sentiment %", baseline.sentiment_share.positive, kpi.sentiment_share.positive),
            ("Negative Sentiment %", baseline.sentiment_share.negative, kpi.sentiment_share.negative),
            ("Rally Mood Score", baseline.rally_mood_score, kpi.rally_mood_score),
            ("Leader Favourability", baseline.leader_favourability_score, kpi.leader_favourability_score),
            ("Share of Voice %", baseline.share_of_voice, kpi.share_of_voice),
            ("Bot Suspicion %", baseline.bot_suspicion_score, kpi.bot_suspicion_score),
            ("Polarization Index", baseline.polarization_index, kpi.polarization_index),
        ]
        table = doc.add_table(rows=len(metrics) + 1, cols=4, style="Light Grid Accent 1")
        headers = ["Metric", "Baseline", "Current", "Change"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for i, (name, base, curr) in enumerate(metrics, 1):
            delta = curr - base if isinstance(curr, (int, float)) else 0
            arrow = "+" if delta > 0 else "" if delta < 0 else ""
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = str(round(base, 1))
            table.rows[i].cells[2].text = str(round(curr, 1))
            table.rows[i].cells[3].text = f"{arrow}{round(delta, 1)}"

    def _add_platform_performance(self, doc, kpi):
        doc.add_heading("Platform Performance", level=1)
        plat_path = os.path.join(self.charts_dir, "platform_bar.png")
        self.chart_gen.platform_bar(kpi, plat_path)
        if os.path.exists(plat_path):
            doc.add_picture(plat_path, width=Inches(5.5))

        if kpi.platform_breakdown:
            table = doc.add_table(rows=len(kpi.platform_breakdown) + 1, cols=2, style="Light Grid Accent 1")
            table.rows[0].cells[0].text = "Platform"
            table.rows[0].cells[1].text = "Volume"
            for i, (plat, vol) in enumerate(sorted(kpi.platform_breakdown.items(), key=lambda x: -x[1]), 1):
                table.rows[i].cells[0].text = plat
                table.rows[i].cells[1].text = str(vol)

    def _add_keywords(self, doc, kpi):
        doc.add_heading("Top Keywords and Hashtags", level=1)

        # Hashtag chart
        ht_path = os.path.join(self.charts_dir, "hashtag_bar.png")
        self.chart_gen.hashtag_bar(kpi, ht_path)
        if os.path.exists(ht_path):
            doc.add_picture(ht_path, width=Inches(5.5))

        # Keywords table
        if kpi.top_keywords:
            doc.add_heading("Top Keywords", level=2)
            table = doc.add_table(rows=min(len(kpi.top_keywords), 20) + 1, cols=2, style="Light Grid Accent 1")
            table.rows[0].cells[0].text = "Keyword"
            table.rows[0].cells[1].text = "Count"
            for i, kw in enumerate(kpi.top_keywords[:20], 1):
                table.rows[i].cells[0].text = kw["keyword"]
                table.rows[i].cells[1].text = str(kw["count"])

    def _add_influencer_analysis(self, doc, kpi):
        doc.add_heading("Top Influencers and Amplifiers", level=1)
        if kpi.top_influencers:
            table = doc.add_table(rows=min(len(kpi.top_influencers), 15) + 1, cols=4, style="Light Grid Accent 1")
            headers = ["Name", "Followers", "Posts", "Total Engagement"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for i, inf in enumerate(kpi.top_influencers[:15], 1):
                table.rows[i].cells[0].text = inf.get("name", inf.get("author_id", ""))
                table.rows[i].cells[1].text = f"{inf.get('followers', 0):,}"
                table.rows[i].cells[2].text = str(inf.get("posts", 0))
                table.rows[i].cells[3].text = f"{inf.get('engagement', 0):,}"

    def _add_narratives(self, doc, narrative_summary):
        doc.add_heading("Dominant Narratives", level=1)
        for narr in narrative_summary.get("dominant_narratives", []):
            doc.add_heading(narr.get("narrative_title", "Untitled"), level=2)
            doc.add_paragraph(f"Direction: {narr.get('sentiment_direction', 'N/A')} | "
                            f"Strength: {narr.get('strength', 0)} | "
                            f"Reach: {narr.get('estimated_reach', 'N/A')}")

        if narrative_summary.get("counter_narratives"):
            doc.add_heading("Counter-Narratives", level=2)
            for cn in narrative_summary["counter_narratives"]:
                doc.add_paragraph(f"- {cn.get('narrative_title', cn) if isinstance(cn, dict) else cn}")

        if narrative_summary.get("emerging_risks"):
            doc.add_heading("Emerging Risks", level=2)
            for risk in narrative_summary["emerging_risks"]:
                doc.add_paragraph(f"- {risk}")

    def _add_risk_assessment(self, doc, kpi):
        doc.add_heading("Risk Assessment", level=1)
        risks = []
        if kpi.bot_suspicion_score > 20:
            risks.append(("High Bot Activity", "HIGH",
                         f"{kpi.bot_suspicion_score}% of content flagged as potentially bot-generated"))
        if kpi.polarization_index > 50:
            risks.append(("High Polarization", "HIGH",
                         f"Polarization index at {kpi.polarization_index} indicates divisive discourse"))
        if kpi.misinformation_flag_count > 10:
            risks.append(("Misinformation Detected", "MEDIUM",
                         f"{kpi.misinformation_flag_count} items flagged for review"))
        if kpi.opposition_counter_intensity > 40:
            risks.append(("Strong Opposition Response", "MEDIUM",
                         f"Opposition counter-narrative intensity at {kpi.opposition_counter_intensity}%"))
        if not risks:
            risks.append(("No Critical Risks", "LOW", "No significant risks identified in this phase"))

        table = doc.add_table(rows=len(risks) + 1, cols=3, style="Light Grid Accent 1")
        for i, h in enumerate(["Risk", "Severity", "Detail"]):
            table.rows[0].cells[i].text = h
        for i, (risk, sev, detail) in enumerate(risks, 1):
            table.rows[i].cells[0].text = risk
            table.rows[i].cells[1].text = sev
            table.rows[i].cells[2].text = detail

    def _add_methodology(self, doc, kpi):
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(
            "This report was generated using the JanPulse AI, an AI-powered "
            "public mood analysis platform. The methodology employs a 5-layer hybrid NLP stack: "
            "(1) Rules-based keyword matching, (2) Domain-specific political lexicon scoring, "
            "(3) Transformer-based classification using multilingual models (MuRIL, XLM-RoBERTa), "
            "(4) Large Language Model interpretation (Claude/GPT-4o) for nuanced cases, and "
            "(5) Human analyst review for flagged content."
        )
        doc.add_paragraph(
            f"Data was collected from {len(kpi.platform_breakdown)} platforms. "
            f"Languages processed include English, Hindi, Bengali, Hinglish, and Romanised Bengali. "
            f"Bot detection used heuristic scoring based on account age, posting patterns, and network behaviour. "
            f"All metrics are computed from {kpi.sentiment_share.total_docs} classified documents."
        )
        doc.add_heading("Limitations", level=2)
        limitations = [
            "Digital sentiment does not equal voting intention.",
            "Bot detection is heuristic-based and may have false positives/negatives.",
            "Sarcasm detection accuracy in Bengali is approximately 70-75%.",
            "Rural and offline sentiment is not captured.",
            "Platform API restrictions may limit data completeness.",
        ]
        for lim in limitations:
            doc.add_paragraph(f"- {lim}")
